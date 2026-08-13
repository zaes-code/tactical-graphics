"""Build the counsel review brief (.docx) from the current plan and memo."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = r'C:\GIT\tactical-graphics\ai\counsel-review-brief.docx'

NAVY = RGBColor(0x0A, 0x16, 0x28)
ACCENT = RGBColor(0x00, 0x47, 0xCC)
MUTED = RGBColor(0x60, 0x6A, 0x78)

doc = Document()

# Base style
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(8)
st.paragraph_format.line_spacing = 1.12

for i, size in ((1, 17), (2, 13), (3, 11.5)):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = NAVY if i > 1 else ACCENT
    h.font.bold = True


def para(text='', bold=False, italic=False, size=None, color=None, space_after=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def rich(parts, style=None, space_after=None):
    """parts: list of (text, bold, italic, mono)."""
    p = doc.add_paragraph(style=style)
    for text, *flags in parts:
        b = flags[0] if len(flags) > 0 else False
        i = flags[1] if len(flags) > 1 else False
        mono = flags[2] if len(flags) > 2 else False
        r = p.add_run(text)
        r.bold, r.italic = b, i
        if mono:
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def quote(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        r.italic = True
    return p


def table(headers, rows, widths=None):
    """headers=None builds a plain key/value table with no header band."""
    ncols = len(headers) if headers else len(rows[0])
    if headers:
        t = doc.add_table(rows=1, cols=ncols)
        t.style = 'Light Grid Accent 1'
        for c, htext in enumerate(headers):
            cell = t.rows[0].cells[c]
            cell.text = ''
            r = cell.paragraphs[0].add_run(htext)
            r.bold = True
            r.font.size = Pt(9.5)
    else:
        t = doc.add_table(rows=0, cols=ncols)
        t.style = 'Light List Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in rows:
        cells = t.add_row().cells
        for c, val in enumerate(row):
            cells[c].text = ''
            r = cells[c].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
            if headers is None and c == 0:
                r.bold = True
    if widths:
        for r_ in t.rows:
            for c, w in enumerate(widths):
                r_.cells[c].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        if isinstance(it, str):
            p.add_run(it)
        else:
            for text, *flags in it:
                b = flags[0] if len(flags) > 0 else False
                mono = flags[1] if len(flags) > 1 else False
                r = p.add_run(text)
                r.bold = b
                if mono:
                    r.font.name = 'Consolas'
                    r.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(4)


# ── Title ───────────────────────────────────────────────────────────────────
para('Counsel Review Brief', bold=True, size=22, color=NAVY, space_after=2)
para('Relicensing @zaes/tactical-graphics from MIT to Business Source License 1.1',
     size=12.5, color=MUTED, space_after=14)

table(
    None,
    [
        ['Prepared for', 'Outside counsel — software licensing and government contracts'],
        ['Prepared by', 'ZAES engineering'],
        ['Date', '13 August 2026'],
        ['Subject', '@zaes/tactical-graphics — a MIL-STD-2525E / FM 1-02.2 symbology library published on npm'],
        ['Decision sought', 'Clearance to relicense forward, and sign-off on two published assertions'],
        ['Status of source documents', 'Both marked DRAFT and not relied upon pending this review'],
    ],
    widths=[1.6, 4.9],
)

para('This brief was prepared internally by engineering. It is not legal advice and '
     'makes no representation that its citations or characterisations are correct — '
     'confirming them is part of what is being asked.',
     italic=True, size=9.5, color=MUTED, space_after=12)

# ── 1. Why ──────────────────────────────────────────────────────────────────
doc.add_heading('1. What we are asking for, in one paragraph', level=1)
rich([
    ('ZAES publishes a tactical-symbology library on npm. Twenty versions have shipped, ', ),
    ('all under the MIT licence', True),
    (', which grants everyone — including competitors and the Government — unlimited '
     'rights in perpetuity, at no cost. We propose to keep those versions exactly as '
     'they are and license ', ),
    ('future major versions', True),
    (' under the Business Source License 1.1: free to install, read, prototype and '
     'evaluate; a paid licence required to run in production; and each release '
     'converting automatically to MIT four years after it ships. We need counsel to '
     'confirm the mechanism works and to clear two statements we intend to publish.', ),
])

# ── 2. Settled ──────────────────────────────────────────────────────────────
doc.add_heading('2. What is already settled, and how', level=1)
para('Two questions that would normally dominate this analysis are closed. Both were '
     'verified against the source repository rather than assumed.', space_after=8)

table(
    ['Question', 'Finding', 'Basis'],
    [
        ['Who owns the copyright?',
         'ZAES, entirely. No third-party interest; no assignment outstanding.',
         'All three contributors were salaried ZAES employees writing within the scope of employment — work made for hire under 17 U.S.C. § 101.'],
        ['Was it developed at private expense?',
         'Yes, exclusively. No CLIN, contract time, or program delivery.',
         'Confirmed by the company; corroborated by a search of the entire commit history and tracked source tree, which contains no contract number, CLIN, CDRL, task order, PWS, SOW or program-office reference.'],
    ],
    widths=[1.5, 2.0, 3.0],
)

para('Two consequences follow, and both are load-bearing for this review:', space_after=6)
bullets([
    [('ZAES may license future versions on any terms it chooses. ', True),
     ('Nothing in the contribution history constrains that.', )],
    [('The software is commercial computer software developed exclusively at private expense, ', True),
     ('so DFARS 252.227-7014 and Government Purpose Rights do not attach, and the Government takes only what the public licence grants.', )],
])

# ── 3. The decisions ────────────────────────────────────────────────────────
doc.add_heading('3. The decisions we need', level=1)
para('Six items gate the relicence. A seventh is needed before the first commercial '
     'quote but not before the licence change, and is flagged as such.', space_after=10)

items = [
    ('3.1', 'Is the Additional Use Grant enforceable, and does it say what we mean?',
     'This clause is the entire commercial model. BUSL supplies a template; the grant is '
     'the part the licensor authors, and it defines the boundary between free use and paid use.',
     'Proposed text (Part 1 of the licensing plan):',
     '“You may use the Licensed Work for any purpose that is not Production Use. '
     '\u2018Production Use\u2019 means using the Licensed Work in, or to support, a system that is '
     'deployed to operational users, delivered to a customer, or used in performance of a '
     'contract. Development, evaluation, testing, prototyping, demonstration, internal '
     'research, and teaching are not Production Use.”',
     'In particular: is “used in performance of a contract” too broad or too narrow for '
     'our market, where the buyer is typically a prime contractor or a government program?'),

    ('3.2', 'Does the output carve-out do its job?',
     'This is a generator library: it takes coordinates and produces GeoJSON that a customer '
     'renders on a map. A buyer\u2019s counsel will ask whether our licence reaches the data '
     'their program produces. We would rather answer that in the licence than in an email.',
     'Proposed text:',
     '“The Licensed Work generates geospatial data. This license governs the Licensed Work '
     'itself; it places no restriction on the GeoJSON or other output your use of it produces.”',
     'Is this sufficient to disclaim any interest in customer output, and does it create any '
     'unintended limitation on our own rights?'),

    ('3.3', 'Are the FAR and DFARS citations correct?',
     'We intend to publish a U.S. Government use block in the package README, where it will be '
     'read by contracting officers. It is a written representation, not marketing.',
     'Proposed text (Part 6 of the licensing plan) asserts the software is commercial computer '
     'software developed exclusively at private expense, and cites:',
     'FAR 12.212, and DFARS 227.7202-1 through 227.7202-4 — commercial computer software is '
     'acquired under the licence customarily provided to the public. It also states that ZAES '
     'asserts restrictions under DFARS 252.227-7017 wherever the software is identified in a proposal.',
     'We believe DFARS 252.227-7014 does not apply (it governs noncommercial software) and that '
     'there is no clause numbered 252.227-7202. Please confirm both, and confirm the block is '
     'safe to publish.'),

    ('3.4', 'Does the private-expense assertion survive indirect cost recovery?',
     'Development was funded from corporate funds. If any of those hours were charged to an '
     'IR&D or B&P account whose costs are recovered as allowable indirect costs on government '
     'contracts, we want to be certain that does not convert the work into government-funded '
     'development for data-rights purposes.',
     'Our position (Provenance Memorandum, § 5):',
     'IR&D and B&P costs recovered as indirect costs remain private expense. Recovering '
     'development cost through an overhead rate does not make the development government-funded.',
     'Please confirm the phrasing against the current DFARS 252.227-7013(a) definitions. This is '
     'a recurring point of confusion between finance and contracts, and we would like it settled '
     'in writing.'),

    ('3.5', 'Is the AI-assisted development position sound, and is disclosure required?',
     'Part of the work was produced with AI coding assistance, directed and reviewed by ZAES '
     'employees. The repository documents that workflow openly, so a diligence reviewer will '
     'find it and ask.',
     'Our position (Provenance Memorandum, § 8):',
     'All AI-assisted output was specified, directed, reviewed, modified and integrated by ZAES '
     'employees, and the resulting work reflects substantial human authorship in selection, '
     'arrangement and modification. This bears on the scope of copyright, not on ownership.',
     'Please confirm the position, and advise whether disclosure is required on the specific '
     'vehicles where this software will be asserted.'),

    ('3.6', 'Does publishing twenty MIT versions impair relicensing forward?',
     'This is the first question a prime\u2019s counsel is likely to raise, so we would like a '
     'clear answer on file before it is asked.',
     'Our understanding:',
     'The MIT grants on versions 1.0.0 through 2.1.0 are irrevocable and perpetual as to those '
     'versions. They do not transfer copyright and do not restrict the licensor\u2019s ability to '
     'license later versions differently. We will not unpublish or obscure them.',
     'Please confirm, and advise whether any notice obligation attaches to the change.'),
]

for num, question, why, lead, text, ask in items:
    doc.add_heading(f'{num}  {question}', level=2)
    rich([('Why it matters. ', True), (why,)], space_after=4)
    rich([(lead,)], space_after=2)
    quote(text)
    rich([('What we need. ', True), (ask,)], space_after=10)

doc.add_heading('3.7  Commercial licence template — needed before the first quote, not before the licence change', level=2)
rich([('Why it matters. ', True),
      ('BUSL is the public default. The negotiated agreement is the enforceable instrument, and '
       'it is where the real protections live: scope of deployment, audit rights, reporting, term '
       'and renewal, assignment on change of control, and remedies.',)], space_after=4)
rich([('What we need. ', True),
      ('A template we can put in front of a prime contractor. ', ),
      ('This does not gate the relicence', True),
      (' — please do not let it hold up items 3.1 to 3.6.',)], space_after=10)

# ── 4. Not asking ───────────────────────────────────────────────────────────
doc.add_heading('4. What we are not asking counsel to decide', level=1)
para('Stated so the review stays scoped, and so nothing here is mistaken for an open question:',
     space_after=6)
bullets([
    [('Whether to relicense at all. ', True), ('That is a business decision and it has been made, subject to this review.',)],
    [('Pricing. ', True), ('Not published anywhere and not part of this review.',)],
    [('The choice of BUSL over AGPL or a proprietary licence. ', True),
     ('Decided on commercial grounds: BUSL states exactly the intended boundary, and AGPL is '
      'rejected outright by many approved-software policies in this market.',)],
    [('Copyright registration mechanics. ', True),
     ('Proceeding in parallel; see § 6.',)],
])

# ── 5. Supporting documents ─────────────────────────────────────────────────
doc.add_heading('5. Supporting documents', level=1)
para('Two internal documents accompany this brief. Both are marked DRAFT and neither is '
     'relied upon pending review.', space_after=8)

table(
    ['Document', 'What it contains', 'Sections for counsel'],
    [
        ['Licensing plan\n(ai/licensing-plan.md)',
         'The relicensing model, the terms, drop-in licence and README copy, the U.S. Government '
         'use block, enforcement approach, and an explanation of where a version\u2019s licence is '
         'actually specified.',
         'Part 1 (model, grant, terms); Part 1b (how per-version licensing works); Part 6 '
         '(government use); Part 7 (enforcement)'],
        ['Provenance memorandum\n(ai/provenance-memo.md)',
         'Ownership and funding analysis, development timeline, per-person contribution record, '
         'third-party components, publication history, and AI-assisted development disclosure. '
         'Every figure derived from the repository is accompanied by the command that reproduces it.',
         '§ 1 (assertion); § 5 (funding and DFARS phrasing); § 8 (AI authorship); § 10 (limitations)'],
    ],
    widths=[1.7, 3.1, 1.7],
)

para('Both remain incomplete in one respect: the HR and finance inputs — employment dates, '
     'IP clauses on file, and charge codes — are marked as placeholders and are being gathered '
     'separately. They do not change the analysis; they are the evidence behind it.',
     space_after=10)

# ── 6. Timing ───────────────────────────────────────────────────────────────
doc.add_heading('6. Timing, and one item that should not wait', level=1)

rich([('Copyright registration is proceeding in parallel and does not depend on this review.', True)],
     space_after=4)
para('Registration is a factual filing about authorship and ownership, both of which are '
     'settled. The identical work would be registered whether it ships under MIT or BUSL. '
     'It is also the only item with a deadline: under 17 U.S.C. § 412, statutory damages and '
     'attorney\u2019s fees are available only if the work was registered before infringement began '
     'or within three months of first publication.', space_after=8)

table(
    ['Version', 'Published', 'Three-month window closes'],
    [
        ['1.0.0', '14 July 2026', '14 October 2026'],
        ['2.0.0', '11 August 2026', '11 November 2026'],
        ['2.1.0', '12 August 2026', '12 November 2026'],
    ],
    widths=[1.2, 2.0, 2.4],
)
para('We would welcome a view on which versions to register and on the deposit approach — '
     'the source is public, so no redaction is expected — but we do not intend to hold the '
     'filing for the licensing opinion.', space_after=10)

doc.add_heading('Version sequencing', level=2)
para('The next major release is intended to be the first BUSL release. It also carries seven '
     'breaking API renames that are ready now. If this review concludes before that release, '
     'both ship together. If it does not, the renames may ship first under MIT and the BUSL '
     'release becomes the following major. Nothing in the analysis depends on which happens.',
     space_after=10)

# ── 7. Background ───────────────────────────────────────────────────────────
doc.add_heading('7. Background counsel may find useful', level=1)

doc.add_heading('The commercial argument, stated plainly', level=2)
para('Under FAR 12.212 the Government takes the licence customarily provided to the public. '
     'Today that licence is MIT — everything, free, in perpetuity, to every agency and to anyone '
     'they pass it to. Under BUSL it becomes a licence requiring a paid grant for production use. '
     'The relicence therefore improves the government-rights position materially, and that is the '
     'strongest single argument for making the change.', space_after=8)

doc.add_heading('One probe we expect', level=2)
para('A contracting officer may ask whether the software is “customarily” licensed to the public '
     'on commercial terms when no commercial licence has yet been sold. FAR 2.101 covers software '
     '“offered for sale, lease, or license to the general public”, which an offer appears to '
     'satisfy — but a view on this would be useful, and the first executed commercial licence '
     'removes the question entirely.', space_after=8)

doc.add_heading('A resolved item, noted for completeness', level=2)
para('Until 12 August 2026 the production dependency tree contained a component licensed under '
     'AGPL-3.0, reached transitively and never used by this library. It was removed in version '
     '2.1.0, and the tree is now 32 components, all permissively licensed and all attributable. '
     'It is recorded here only because it was present in versions 1.0.0 through 2.0.0, which '
     'remain published — if counsel sees any exposure arising from that historical period, we '
     'would like to know.', space_after=10)

# ── Appendix ────────────────────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('Appendix — proposed licence parameters', level=1)
para('BUSL-1.1 is a template with five licensor-supplied fields. Ours would read:', space_after=6)

table(
    ['Field', 'Proposed value'],
    [
        ['Licensor', 'ZAES'],
        ['Licensed Work', '@zaes/tactical-graphics, the first major version published after this review, and later'],
        ['Additional Use Grant', 'As quoted at § 3.1 and § 3.2'],
        ['Change Date', 'Each version\u2019s publication date plus four years (the maximum BUSL permits)'],
        ['Change License', 'MIT'],
    ],
    widths=[1.6, 4.9],
)

para('Each published version carries its own Change Date, so releases convert to MIT on a '
     'rolling schedule, each on its own clock. The conversion is written into the licence text '
     'shipped inside each release, so it cannot be withdrawn — the licensee already holds the file.',
     space_after=8)

para('Three points of mechanism that counsel may wish to confirm are correctly understood:', space_after=6)
bullets([
    [('A version\u2019s licence lives inside its own published archive. ', True),
     ('There is no central registry. Each npm release is immutable and carries its own LICENSE '
      'file and licence metadata, which is why the MIT versions cannot be altered.',)],
    [('A licence change is a breaking change ', True),
     ('and will be published only in a major version, so that no existing user is upgraded into '
      'restricted terms automatically.',)],
    [('BUSL is source-available, not open source. ', True),
     ('It is not OSI-approved, and we will not describe it as open source in any public material.',)],
])

doc.save(OUT)
print('written:', OUT)
